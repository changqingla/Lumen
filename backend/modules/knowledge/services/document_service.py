"""Document service business logic."""

import hashlib
import re
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from fastapi import HTTPException, status
from modules.knowledge.repositories.kb_repository import KnowledgeBaseRepository
from modules.knowledge.repositories.document_repository import DocumentRepository
from repositories.user_repository import UserRepository
from modules.organization.repositories.organization_member_repository import (
    OrganizationMemberRepository,
)
from utils.minio_client import (
    delete_file,
    get_object_metadata,
    get_upload_url,
    temporary_download,
    upload_file,
    upload_file_from_path,
)
from utils.document_process_service import (
    DOCUMENT_ERROR_MESSAGES,
    DocumentProcessService,
    DocumentProcessingError,
    public_document_error,
    sanitize_persisted_document_error,
)
from utils.mineru_service import MineruService
from utils.es_utils import get_user_es_index
from modules.knowledge.entities.document import Document
from modules.knowledge.entities.knowledge_base import KnowledgeBase
from modules.knowledge.document_task_queue import (
    cancel_document_task,
    enqueue_document_task,
)
from config.settings import settings
from typing import Iterator, List, Tuple, Optional
from contextlib import contextmanager
import os
import logging
import asyncio
from pathlib import Path
import tempfile
import uuid

logger = logging.getLogger(__name__)
_OPAQUE_LOG_ID_RE = re.compile(r"^[A-Za-z0-9_-]{1,128}$")


def _opaque_log_id(value: object) -> str:
    candidate = str(value or "").strip()
    return candidate if _OPAQUE_LOG_ID_RE.fullmatch(candidate) else "invalid"


def _log_document_failure(
    *,
    stage: str,
    exc: BaseException,
    document_id: object | None = None,
    task_id: object | None = None,
    level: int = logging.ERROR,
) -> None:
    error_type = (
        exc.error_type
        if isinstance(exc, DocumentProcessingError)
        else type(exc).__name__
    )
    logger.log(
        level,
        "document_processing stage=%s document_id=%s task_id=%s error_type=%s",
        stage,
        _opaque_log_id(document_id),
        _opaque_log_id(task_id),
        re.sub(r"[^A-Za-z0-9_.-]+", "_", error_type)[:80] or "Error",
    )


class DocumentService:
    """Service for document operations."""

    # Supported file extensions (知识库只支持这5种格式)
    PDF_EXTENSIONS = {".pdf"}
    TEXT_EXTENSIONS = {".md", ".markdown", ".txt"}
    WORD_EXTENSIONS = {".doc", ".docx"}
    RAG_TERMINAL_TASK_STATUSES = frozenset({"cancelled", "failed", "completed"})
    TEXT_READ_CHUNK_SIZE = 64 * 1024
    MARKDOWN_DOWNLOAD_CONCURRENCY = 4

    def __init__(self, db: AsyncSession):
        self.db = db
        self.kb_repo = KnowledgeBaseRepository(db)
        self.doc_repo = DocumentRepository(db)
        self.user_repo = UserRepository(db)
        self.org_member_repo = OrganizationMemberRepository(db)

    async def _verify_kb_write_access(self, kb_id: str, user_id: str) -> KnowledgeBase:
        """
        Verify user has WRITE access to knowledge base.
        Only owner and admin users have write permissions.

        Returns:
            Knowledge base object if user has write access

        Raises:
            HTTPException: If knowledge base not found or user has no write permission
        """
        user = await self.user_repo.get_by_id(uuid.UUID(user_id))
        kb = await self.kb_repo.get_writable_by_id(
            kb_id,
            user_id,
            is_admin=bool(user and user.is_admin),
        )

        if not kb:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail={
                    "error": {
                        "code": "FORBIDDEN",
                        "message": "Only the knowledge base owner or admin can perform this action",
                    }
                },
            )

        return kb

    async def _verify_kb_access(self, kb_id: str, user_id: str) -> KnowledgeBase:
        """
        Verify user has access to knowledge base.
        - Admin users: can access any knowledge base
        - Owners: can access their own knowledge bases
        - Organization members: can access organization-shared knowledge bases
        - Everyone: can access public knowledge bases

        Returns:
            Knowledge base object if accessible

        Raises:
            HTTPException: If knowledge base not found or not accessible
        """
        user_uuid = uuid.UUID(user_id)
        user = await self.user_repo.get_by_id(user_uuid)
        is_admin = bool(user and user.is_admin)
        user_org_ids = (
            []
            if is_admin
            else await self.org_member_repo.get_user_org_ids(user_uuid)
        )
        kb = await self.kb_repo.get_accessible_by_id(
            kb_id,
            user_uuid,
            user_org_ids,
            is_admin=is_admin,
        )
        if not kb:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={
                    "error": {
                        "code": "NOT_FOUND",
                        "message": "Knowledge base not found or not accessible",
                    }
                },
            )

        return kb

    def _get_file_extension(self, filename: str) -> str:
        """Get file extension in lowercase."""
        return os.path.splitext(filename)[1].lower()

    def _normalize_filename(self, filename: Optional[str]) -> str:
        """Normalize filename and remove path traversal segments."""
        safe_name = os.path.basename(filename or "").strip()
        if not safe_name:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "INVALID_REQUEST",
                        "message": "filename is required",
                    }
                },
            )
        return safe_name

    def _build_object_name(
        self, storage_owner_id: str, kb_id: str, filename: str
    ) -> str:
        """Build a unique object name to avoid overwrite on duplicate filenames."""
        return f"kb/{storage_owner_id}/{kb_id}/{uuid.uuid4().hex}_{filename}"

    @staticmethod
    def _get_owner_es_index(kb: KnowledgeBase) -> str:
        """Resolve the ES index from the knowledge base owner, not the acting user."""
        return get_user_es_index(str(kb.owner_id))

    def _needs_mineru_conversion(self, filename: str) -> bool:
        """Check if file needs Mineru conversion (PDF only)."""
        ext = self._get_file_extension(filename)
        return ext in self.PDF_EXTENSIONS

    async def _enqueue_persisted_document(self, document_id: str) -> None:
        """Best-effort enqueue; the queue reconciler repairs Redis outages."""
        try:
            await enqueue_document_task(document_id)
        except Exception as exc:
            _log_document_failure(
                stage="queue_enqueue",
                document_id=document_id,
                exc=exc,
            )

    def _extract_docx_content_from_path(self, file_path: str | os.PathLike[str]) -> str:
        """Extract DOCX content using python-docx's file-path API."""
        from docx import Document as DocxDocument

        return self._extract_docx_text(DocxDocument(str(file_path)))

    @staticmethod
    def _extract_docx_text(doc) -> str:
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        content = "\n\n".join(paragraphs)
        if not content:
            raise ValueError("DOCX file contains no extractable text")
        return content

    def _extract_doc_content_from_path(self, file_path: str | os.PathLike[str]) -> str:
        """Extract legacy DOC content using Tika's file-path API."""
        from tika import parser

        result = parser.from_file(str(file_path))
        content = result.get("content", "")
        if not content:
            return ""
        lines = [line.strip() for line in content.split("\n") if line.strip()]
        return "\n\n".join(lines)

    @classmethod
    def _read_text_file(cls, file_path: str | os.PathLike[str]) -> str:
        """Decode a text document incrementally with the legacy fallbacks."""
        for encoding in ("utf-8", "gbk", "latin-1"):
            chunks: list[str] = []
            try:
                with Path(file_path).open("r", encoding=encoding) as source:
                    while True:
                        chunk = source.read(cls.TEXT_READ_CHUNK_SIZE)
                        if not chunk:
                            break
                        chunks.append(chunk)
            except UnicodeDecodeError:
                continue
            return "".join(chunks)
        raise UnicodeDecodeError("utf-8", b"", 0, 1, "Unable to decode document")

    @classmethod
    @contextmanager
    def _temporary_markdown_file(
        cls,
        markdown_content: str,
        doc_id: str,
    ) -> Iterator[Path]:
        """Write Markdown in bounded chunks and remove it on every exit path."""
        temp_path: Path | None = None
        try:
            with tempfile.NamedTemporaryFile(
                mode="w",
                encoding="utf-8",
                prefix=f"lumen-document-{doc_id}-",
                suffix=".md",
                delete=False,
            ) as output:
                temp_path = Path(output.name)
                for offset in range(0, len(markdown_content), cls.TEXT_READ_CHUNK_SIZE):
                    output.write(
                        markdown_content[offset : offset + cls.TEXT_READ_CHUNK_SIZE]
                    )
            yield temp_path
        finally:
            if temp_path is not None:
                temp_path.unlink(missing_ok=True)

    async def init_direct_upload(
        self,
        kb_id: str,
        user_id: str,
        filename: str,
        file_size: int,
        content_type: Optional[str] = None,
    ) -> dict:
        """
        Initialize direct browser upload to MinIO.
        Returns upload URL + pre-created document record.
        """
        kb = await self._verify_kb_write_access(kb_id, user_id)

        safe_filename = self._normalize_filename(filename)
        ext = self._get_file_extension(safe_filename)
        supported_extensions = (
            self.PDF_EXTENSIONS | self.TEXT_EXTENSIONS | self.WORD_EXTENSIONS
        )
        if ext not in supported_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "UNSUPPORTED_FORMAT",
                        "message": f"Unsupported file format: {ext}. Supported formats: pdf, txt, md, doc, docx",
                    }
                },
            )

        if file_size <= 0:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {"code": "INVALID_REQUEST", "message": "Invalid file size"}
                },
            )

        if file_size > settings.MAX_UPLOAD_SIZE:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "FILE_TOO_LARGE",
                        "message": f"File too large. Max size is {settings.MAX_UPLOAD_SIZE // 1024 // 1024}MB",
                    }
                },
            )

        storage_owner_id = str(kb.owner_id)
        object_name = self._build_object_name(storage_owner_id, kb_id, safe_filename)
        upload_url = get_upload_url(object_name, expires_seconds=900)

        document = await self.doc_repo.create(
            kb_id=kb_id,
            name=safe_filename,
            size=file_size,
            source="upload",
            file_path=f"{settings.MINIO_BUCKET}/{object_name}",
        )

        return {
            "id": str(document.id),
            "name": document.name,
            "status": document.status,
            "uploadUrl": upload_url,
        }

    async def create_markdown_document_from_content(
        self,
        *,
        kb_id: str,
        user_id: str,
        filename: str,
        markdown: str,
        source: str,
    ) -> Document:
        """Create or reuse a Markdown document from in-memory content."""
        kb = await self._verify_kb_write_access(kb_id, user_id)
        existing = await self.doc_repo.get_by_kb_and_source(kb_id, source)
        if existing:
            return existing

        safe_filename = self._normalize_filename(filename)
        if self._get_file_extension(safe_filename) not in self.TEXT_EXTENSIONS:
            safe_filename = f"{os.path.splitext(safe_filename)[0] or 'document'}.md"

        file_data = markdown.encode("utf-8")
        storage_owner_id = str(kb.owner_id)
        object_name = self._build_object_name(storage_owner_id, kb_id, safe_filename)
        file_path = await upload_file(object_name, file_data, "text/markdown")

        document = await self.doc_repo.create(
            kb_id=kb_id,
            name=safe_filename,
            size=len(file_data),
            source=source,
            file_path=file_path,
        )

        await self.doc_repo.update_status(
            document,
            Document.STATUS_QUEUED,
            error_message=None,
        )
        await self._enqueue_persisted_document(str(document.id))
        return document

    async def complete_direct_upload(
        self,
        kb_id: str,
        user_id: str,
        doc_id: str,
    ) -> dict:
        """
        Complete direct upload and start background processing.
        """
        await self._verify_kb_write_access(kb_id, user_id)
        doc = await self.doc_repo.get_by_id(doc_id, kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={
                    "error": {"code": "NOT_FOUND", "message": "Document not found"}
                },
            )

        if doc.status == Document.STATUS_QUEUED:
            await self._enqueue_persisted_document(str(doc.id))
            return {
                "id": str(doc.id),
                "name": doc.name,
                "status": Document.STATUS_PROCESSING,
            }

        if doc.status != Document.STATUS_UPLOADING:
            return {
                "id": str(doc.id),
                "name": doc.name,
                "status": Document.public_status(doc.status),
            }

        if not doc.file_path:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "INVALID_STATE",
                        "message": "Document file path is missing",
                    }
                },
            )

        object_name = doc.file_path.replace(f"{settings.MINIO_BUCKET}/", "")
        object_metadata = await get_object_metadata(object_name)
        if object_metadata is None:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "UPLOAD_NOT_FINISHED",
                        "message": "File not uploaded yet",
                    }
                },
            )
        if object_metadata.size != doc.size:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "UPLOAD_SIZE_MISMATCH",
                        "message": "Uploaded file size does not match the initialized upload",
                    }
                },
            )
        if object_metadata.size <= 0 or object_metadata.size > settings.MAX_UPLOAD_SIZE:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "INVALID_UPLOAD_SIZE",
                        "message": "Uploaded file size is outside the allowed range",
                    }
                },
            )

        await self.doc_repo.update_status(
            doc,
            Document.STATUS_QUEUED,
            error_message=None,
        )
        await self._enqueue_persisted_document(str(doc.id))
        logger.info("document_processing stage=queued document_id=%s", doc.id)

        return {
            "id": str(doc.id),
            "name": doc.name,
            "status": Document.STATUS_PROCESSING,
        }

    async def process_queued_document(self, doc_id: str) -> str:
        """Resolve and process a durable queue item from authoritative state.

        The method is intentionally resumable at the two externally durable
        polling stages. Other interrupted stages restart from the persisted
        Markdown or original object after removing stale search data.
        """
        result = await self.db.execute(select(Document).where(Document.id == doc_id))
        doc = result.scalar_one_or_none()
        if doc is None:
            return "missing"

        kb = await self.kb_repo.get_by_id_any(str(doc.kb_id))
        if kb is None:
            await self.doc_repo.update_status(
                doc,
                Document.STATUS_FAILED,
                error_message=DOCUMENT_ERROR_MESSAGES["knowledge_base"],
            )
            return Document.STATUS_FAILED

        if doc.status == Document.STATUS_READY:
            await self.kb_repo.sync_contents_count(doc.kb_id)
            return Document.STATUS_READY

        es_index_name = self._get_owner_es_index(kb)

        # A RAG task already has its own durable queue, so resume polling it
        # rather than submit a duplicate after this API worker restarts.
        if doc.status == Document.STATUS_EMBEDDING and doc.parse_task_id:
            success = await self._poll_parse_task(doc, doc.parse_task_id, self.doc_repo)
            if success:
                await self.kb_repo.sync_contents_count(doc.kb_id)
            return await self._reload_document_status(doc_id)

        # MinerU task identifiers are durable too. Resume conversion polling
        # and persist the result before entering the chunking stage.
        if (
            doc.status == Document.STATUS_PROCESSING
            and doc.mineru_task_id
            and self._needs_mineru_conversion(doc.name)
        ):
            markdown_content = await self._poll_mineru_task(doc.mineru_task_id, doc_id)
            md_object_name = f"kb/{kb.owner_id}/{doc.kb_id}/markdown/{doc_id}.md"
            try:
                with self._temporary_markdown_file(
                    markdown_content, doc_id
                ) as temp_path:
                    markdown_path = await upload_file_from_path(
                        md_object_name,
                        temp_path,
                        "text/markdown",
                    )
            except Exception as exc:
                raise DocumentProcessingError(
                    "storage",
                    error_type=type(exc).__name__,
                ) from None
            await self.doc_repo.update_markdown_path(
                doc,
                markdown_path,
                hashlib.sha256(markdown_content.encode("utf-8")).hexdigest(),
            )
            await self._delete_existing_document_index(doc_id, es_index_name)
            await self._retry_chunking_only(
                doc_id,
                es_index_name,
                markdown_content,
                doc.name,
            )
            return await self._reload_document_status(doc_id)

        await self.doc_repo.update_status(
            doc,
            Document.STATUS_PROCESSING,
            error_message=None,
        )

        if doc.markdown_path:
            try:
                markdown_object_name = doc.markdown_path.replace(
                    f"{settings.MINIO_BUCKET}/",
                    "",
                )
                async with temporary_download(
                    markdown_object_name,
                    suffix=".md",
                    max_bytes=settings.MAX_UPLOAD_SIZE,
                ) as markdown_file_path:
                    markdown_content = await asyncio.to_thread(
                        self._read_text_file,
                        markdown_file_path,
                    )
                if markdown_content:
                    actual_sha256 = hashlib.sha256(
                        markdown_content.encode("utf-8")
                    ).hexdigest()
                    expected_sha256 = (
                        str(getattr(doc, "markdown_sha256", None) or "").strip().lower()
                    )
                    if expected_sha256 and actual_sha256 != expected_sha256:
                        raise RuntimeError(
                            "Persisted Markdown failed its SHA-256 integrity check"
                        )
                    if not expected_sha256:
                        await self.doc_repo.update_markdown_path(
                            doc,
                            doc.markdown_path,
                            actual_sha256,
                        )
                    await self._delete_existing_document_index(doc_id, es_index_name)
                    await self._retry_chunking_only(
                        doc_id,
                        es_index_name,
                        markdown_content,
                        doc.name,
                    )
                    return await self._reload_document_status(doc_id)
            except Exception as exc:
                _log_document_failure(
                    stage="persisted_markdown_reuse",
                    document_id=doc_id,
                    exc=exc,
                    level=logging.WARNING,
                )

        if not doc.file_path:
            await self.doc_repo.update_status(
                doc,
                Document.STATUS_FAILED,
                error_message=DOCUMENT_ERROR_MESSAGES["source"],
            )
            return Document.STATUS_FAILED

        await self._delete_existing_document_index(doc_id, es_index_name)
        object_name = doc.file_path.replace(f"{settings.MINIO_BUCKET}/", "")
        try:
            async with temporary_download(
                object_name,
                suffix=self._get_file_extension(doc.name),
                max_bytes=settings.MAX_UPLOAD_SIZE,
            ) as file_path:
                if file_path.stat().st_size == 0:
                    raise DocumentProcessingError(
                        "empty",
                        error_type="EmptyDocument",
                    )
                await self._process_document_pipeline(
                    doc_id,
                    es_index_name,
                    file_path,
                    doc.name,
                )
        except DocumentProcessingError:
            raise
        except Exception as exc:
            raise DocumentProcessingError(
                "source",
                error_type=type(exc).__name__,
            ) from None
        return await self._reload_document_status(doc_id)

    async def _reload_document_status(self, doc_id: str) -> str:
        result = await self.db.execute(
            select(Document)
            .where(Document.id == doc_id)
            .execution_options(populate_existing=True)
        )
        doc = result.scalar_one_or_none()
        return doc.status if doc is not None else "missing"

    @staticmethod
    async def _delete_existing_document_index(doc_id: str, es_index_name: str) -> None:
        try:
            await DocumentProcessService.delete_document_from_es(doc_id, es_index_name)
        except Exception as exc:
            _log_document_failure(
                stage="search_cleanup",
                document_id=doc_id,
                exc=exc,
                level=logging.WARNING,
            )

    async def _settle_rag_task_for_deletion(self, task_id: str) -> bool:
        """Cancel an active RAG task and wait briefly for a write-safe terminal state."""
        task_status = await DocumentProcessService.get_task_status(task_id)
        normalized_status = str(task_status.get("status") or "").lower()
        if normalized_status in self.RAG_TERMINAL_TASK_STATUSES:
            return True

        cancellation = await DocumentProcessService.cancel_task(task_id)
        cancellation_state = str(cancellation.get("state") or "").lower()
        cancelled_task_status = str(
            (cancellation.get("task") or {}).get("status") or ""
        ).lower()
        if (
            cancellation_state == "cancelled"
            or cancelled_task_status in self.RAG_TERMINAL_TASK_STATUSES
        ):
            return True
        if cancellation_state != "cancellation_requested":
            return False

        loop = asyncio.get_running_loop()
        deadline = loop.time() + max(
            float(settings.KNOWLEDGE_DOCUMENT_RAG_CANCEL_WAIT_SECONDS),
            0.0,
        )
        while loop.time() < deadline:
            await asyncio.sleep(min(0.5, max(deadline - loop.time(), 0.0)))
            task_status = await DocumentProcessService.get_task_status(task_id)
            normalized_status = str(task_status.get("status") or "").lower()
            if normalized_status in self.RAG_TERMINAL_TASK_STATUSES:
                return True
        return False

    async def _process_document_pipeline(
        self,
        doc_id: str,
        es_index_name: str,
        file_path: str | os.PathLike[str],
        filename: str,
    ):
        """
        Background task to process document through complete pipeline.

        Pipeline:
        1. Convert with Mineru (if PDF)
        2. Parse document (chunk + embed + store)
        3. Update status
        """
        # Import here to avoid circular dependency
        from config.database import AsyncSessionLocal
        from modules.knowledge.repositories.document_repository import (
            DocumentRepository,
        )
        from modules.knowledge.repositories.kb_repository import KnowledgeBaseRepository

        # Create new DB session for background task
        async with AsyncSessionLocal() as db:
            doc_repo = DocumentRepository(db)
            kb_repo = KnowledgeBaseRepository(db)

            # Get document
            result = await db.execute(select(Document).where(Document.id == doc_id))
            doc = result.scalar_one_or_none()
            if not doc:
                logger.error(
                    "document_processing stage=load document_id=%s task_id=none error_type=DocumentNotFound",
                    _opaque_log_id(doc_id),
                )
                return

            logger.info(
                "document_processing stage=pipeline_start document_id=%s",
                _opaque_log_id(doc_id),
            )

            try:
                markdown_content = None
                ext = self._get_file_extension(filename)

                # Step 1: Extract content based on file type
                if self._needs_mineru_conversion(filename):
                    # PDF files: use MinerU for conversion
                    logger.info(
                        "document_processing stage=conversion_submit document_id=%s",
                        _opaque_log_id(doc_id),
                    )
                    await doc_repo.update_status(doc, Document.STATUS_PROCESSING)

                    try:
                        mineru_result = await MineruService.convert_document_from_path(
                            file_path,
                            filename,
                        )
                        batch_id = mineru_result["batch_id"]
                        await doc_repo.update_status(
                            doc, Document.STATUS_PROCESSING, mineru_task_id=batch_id
                        )

                        markdown_content = await self._poll_mineru_task(
                            batch_id, doc_id
                        )
                        logger.info(
                            "document_processing stage=conversion_complete document_id=%s task_id=%s",
                            _opaque_log_id(doc_id),
                            _opaque_log_id(batch_id),
                        )

                    except Exception as exc:
                        _log_document_failure(
                            stage="conversion",
                            document_id=doc_id,
                            task_id=locals().get("batch_id"),
                            exc=exc,
                        )
                        await doc_repo.update_status(
                            doc,
                            Document.STATUS_FAILED,
                            error_message=DOCUMENT_ERROR_MESSAGES["conversion"],
                        )
                        return

                elif ext == ".docx":
                    logger.info(
                        "document_processing stage=extraction document_id=%s",
                        _opaque_log_id(doc_id),
                    )
                    await doc_repo.update_status(doc, Document.STATUS_PROCESSING)
                    try:
                        markdown_content = await asyncio.to_thread(
                            self._extract_docx_content_from_path,
                            file_path,
                        )
                    except Exception as exc:
                        _log_document_failure(
                            stage="extraction",
                            document_id=doc_id,
                            exc=exc,
                        )
                        await doc_repo.update_status(
                            doc,
                            Document.STATUS_FAILED,
                            error_message=DOCUMENT_ERROR_MESSAGES["extraction"],
                        )
                        return

                elif ext == ".doc":
                    # DOC files: use Apache Tika to extract content
                    logger.info(
                        "document_processing stage=extraction document_id=%s",
                        _opaque_log_id(doc_id),
                    )
                    await doc_repo.update_status(doc, Document.STATUS_PROCESSING)
                    try:
                        markdown_content = await asyncio.to_thread(
                            self._extract_doc_content_from_path,
                            file_path,
                        )
                        if not markdown_content:
                            raise Exception("Tika returned empty content")
                    except Exception as exc:
                        _log_document_failure(
                            stage="extraction",
                            document_id=doc_id,
                            exc=exc,
                        )
                        await doc_repo.update_status(
                            doc,
                            Document.STATUS_FAILED,
                            error_message=DOCUMENT_ERROR_MESSAGES["extraction"],
                        )
                        return

                else:
                    # Text files (.txt, .md, .markdown): read directly
                    logger.info(
                        "document_processing stage=extraction document_id=%s",
                        _opaque_log_id(doc_id),
                    )
                    try:
                        markdown_content = await asyncio.to_thread(
                            self._read_text_file,
                            file_path,
                        )
                    except Exception as exc:
                        raise DocumentProcessingError(
                            "extraction",
                            error_type=type(exc).__name__,
                        ) from None

                # Save markdown content to MinIO (for agent use)
                kb_id = str(doc.kb_id)
                kb = await kb_repo.get_by_id_any(kb_id)
                if kb is None:
                    raise DocumentProcessingError(
                        "knowledge_base",
                        error_type="KnowledgeBaseNotFound",
                    )
                storage_owner_id = str(kb.owner_id)
                md_object_name = f"kb/{storage_owner_id}/{kb_id}/markdown/{doc_id}.md"

                # Step 2: Parse document (chunk + embed + store to ES)
                await doc_repo.update_status(doc, Document.STATUS_CHUNKING)
                parse_filename = os.path.splitext(filename)[0] + ".md"
                with self._temporary_markdown_file(
                    markdown_content, doc_id
                ) as temp_file_path:
                    try:
                        markdown_path = await upload_file_from_path(
                            md_object_name,
                            temp_file_path,
                            "text/markdown",
                        )
                        await doc_repo.update_markdown_path(
                            doc,
                            markdown_path,
                            hashlib.sha256(
                                markdown_content.encode("utf-8")
                            ).hexdigest(),
                        )
                    except Exception as exc:
                        raise DocumentProcessingError(
                            "storage",
                            error_type=type(exc).__name__,
                        ) from None
                    logger.info(
                        "document_processing stage=index_submit document_id=%s",
                        _opaque_log_id(doc_id),
                    )
                    parse_result = await DocumentProcessService.parse_document(
                        str(temp_file_path), str(doc_id), es_index_name, parse_filename
                    )

                task_id = parse_result["task_id"]
                logger.info(
                    "document_processing stage=index_accepted document_id=%s task_id=%s",
                    _opaque_log_id(doc_id),
                    _opaque_log_id(task_id),
                )

                await doc_repo.update_status(
                    doc, Document.STATUS_EMBEDDING, parse_task_id=task_id
                )

                success = await self._poll_parse_task(doc, task_id, doc_repo)

                if success:
                    logger.info(
                        "document_processing stage=complete document_id=%s task_id=%s",
                        _opaque_log_id(doc_id),
                        _opaque_log_id(task_id),
                    )
                    # Increment KB contents count only on success
                    await kb_repo.sync_contents_count(doc.kb_id)

            except Exception as exc:
                _log_document_failure(
                    stage=(
                        exc.stage
                        if isinstance(exc, DocumentProcessingError)
                        else "processing"
                    ),
                    document_id=doc_id,
                    exc=exc,
                )
                await doc_repo.update_status(
                    doc,
                    Document.STATUS_FAILED,
                    error_message=public_document_error(exc),
                )

    async def _poll_mineru_task(
        self, batch_id: str, doc_id: str = None, max_attempts: int = 180
    ) -> str:
        """
        Poll MinerU task until completion.

        Args:
            batch_id: MinerU batch ID
            doc_id: Document ID for logging
            max_attempts: Maximum polling attempts (default: 180 = 15 minutes with 5s interval)

        Returns:
            Markdown content string
        """
        for _ in range(max_attempts):
            await asyncio.sleep(5)  # Wait 5 seconds

            try:
                task_status = await MineruService.get_task_status(batch_id)
                status_value = task_status.get("status")
                status = (
                    status_value.strip().lower()
                    if isinstance(status_value, str)
                    else ""
                )
                if status == "completed":
                    return await MineruService.get_content(batch_id)
                if status in {"failed", "cancelled"}:
                    failure = DocumentProcessingError(
                        "conversion",
                        error_type="RemoteTaskFailed",
                    )
                    _log_document_failure(
                        stage="conversion",
                        document_id=doc_id,
                        task_id=batch_id,
                        exc=failure,
                    )
                    raise failure
                if status not in {"pending", "running"}:
                    raise ValueError("invalid conversion task status")

            except DocumentProcessingError:
                raise
            except Exception as exc:
                _log_document_failure(
                    stage="conversion_poll",
                    document_id=doc_id,
                    task_id=batch_id,
                    exc=exc,
                    level=logging.WARNING,
                )

        raise DocumentProcessingError(
            "conversion",
            error_type="TaskTimeout",
        )

    async def _poll_parse_task(
        self, doc: Document, task_id: str, doc_repo, max_attempts: int = 120
    ) -> bool:
        """
        Poll document parsing task until completion (timeout: 10 minutes).

        Returns:
            True if processing succeeded, False if failed
        """
        for _ in range(max_attempts):
            await asyncio.sleep(5)  # Wait 5 seconds

            try:
                task_status = await DocumentProcessService.get_task_status(task_id)
                status_value = task_status.get("status")
                status = (
                    status_value.strip().lower()
                    if isinstance(status_value, str)
                    else ""
                )

                if status == "completed":
                    # Get chunk count from task data
                    chunk_count = task_status.get("data", {}).get("total_chunks", 0)
                    await doc_repo.update_status(
                        doc, Document.STATUS_READY, chunk_count=chunk_count
                    )
                    logger.info(
                        "document_processing stage=index_complete document_id=%s task_id=%s",
                        _opaque_log_id(doc.id),
                        _opaque_log_id(task_id),
                    )
                    return True

                if status == "failed":
                    failure = DocumentProcessingError(
                        "index_failed",
                        error_type="RemoteTaskFailed",
                    )
                    _log_document_failure(
                        stage="index_failed",
                        document_id=doc.id,
                        task_id=task_id,
                        exc=failure,
                    )
                    await doc_repo.update_status(
                        doc,
                        Document.STATUS_FAILED,
                        error_message=DOCUMENT_ERROR_MESSAGES["index_failed"],
                    )
                    return False
                if status not in {
                    "queued",
                    "pending",
                    "processing",
                    "chunking",
                    "embedding",
                    "storing",
                    "cancelled",
                }:
                    raise ValueError("invalid indexing task status")

            except Exception as exc:
                _log_document_failure(
                    stage="index_status",
                    document_id=doc.id,
                    task_id=task_id,
                    exc=exc,
                    level=logging.WARNING,
                )

        await doc_repo.update_status(
            doc,
            Document.STATUS_FAILED,
            error_message=DOCUMENT_ERROR_MESSAGES["index_timeout"],
        )
        return False

    async def list_documents(
        self, kb_id: str, user_id: str, page: int = 1, page_size: int = 20
    ) -> Tuple[List[dict], int]:
        """List documents in knowledge base (admin users can access any KB)."""
        # Verify access permission
        await self._verify_kb_access(kb_id, user_id)

        documents, total = await self.doc_repo.list_documents(kb_id, page, page_size)
        return [doc.to_dict() for doc in documents], total

    async def list_materialized_document_ids(
        self,
        kb_id: str,
        user_id: str,
        page: int = 1,
        page_size: int = 100,
    ) -> Tuple[List[str], int]:
        """List accessible documents with a committed Runtime Markdown revision."""
        await self._verify_kb_access(kb_id, user_id)
        return await self.doc_repo.list_materialized_document_ids(
            kb_id,
            page,
            page_size,
        )

    async def get_document_status(self, doc_id: str, kb_id: str, user_id: str) -> dict:
        """Get document processing status (admin users can access any KB)."""
        # Verify access permission
        await self._verify_kb_access(kb_id, user_id)

        doc = await self.doc_repo.get_by_id(doc_id, kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={
                    "error": {"code": "NOT_FOUND", "message": "Document not found"}
                },
            )

        return {
            "status": Document.public_status(doc.status),
            "errorMessage": sanitize_persisted_document_error(doc.error_message),
            "chunkCount": doc.chunk_count,
        }

    async def get_document_url(self, doc_id: str, kb_id: str, user_id: str) -> dict:
        """Get presigned URL for document file (admin users can access any KB)."""
        from utils.minio_client import get_file_url

        # Verify access permission
        await self._verify_kb_access(kb_id, user_id)

        doc = await self.doc_repo.get_by_id(doc_id, kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={
                    "error": {"code": "NOT_FOUND", "message": "Document not found"}
                },
            )

        if not doc.file_path:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={
                    "error": {"code": "NOT_FOUND", "message": "Document file not found"}
                },
            )

        # Extract object name from file_path
        object_name = doc.file_path.replace(f"{settings.MINIO_BUCKET}/", "")

        # Generate presigned URL (valid for 1 hour)
        file_url = get_file_url(object_name, expires_seconds=3600)

        return {"url": file_url, "name": doc.name}

    async def get_document_markdown(self, doc_id: str, kb_id: str, user_id: str) -> str:
        """Get markdown content of a document (admin users can access any KB)."""
        from utils.minio_client import download_file

        # Verify access permission
        await self._verify_kb_access(kb_id, user_id)

        doc = await self.doc_repo.get_by_id(doc_id, kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={
                    "error": {"code": "NOT_FOUND", "message": "Document not found"}
                },
            )

        if not doc.markdown_path:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={
                    "error": {
                        "code": "NOT_FOUND",
                        "message": "Markdown content not available",
                    }
                },
            )

        # Extract object name from markdown_path
        object_name = doc.markdown_path.replace(f"{settings.MINIO_BUCKET}/", "")

        try:
            # Download markdown from MinIO
            markdown_bytes = await download_file(object_name)
            markdown_content = markdown_bytes.decode("utf-8")
            return markdown_content
        except Exception as exc:
            _log_document_failure(
                stage="markdown_download",
                document_id=doc_id,
                exc=exc,
            )
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail={
                    "error": {
                        "code": "INTERNAL_ERROR",
                        "message": "Failed to retrieve markdown content",
                    }
                },
            ) from None

    async def get_documents_markdown_batch(
        self, doc_ids: List[str], kb_id: str, user_id: str
    ) -> dict:
        """
        Batch get markdown content of multiple documents (for agent use).

        Args:
            doc_ids: List of document IDs
            kb_id: Knowledge base ID
            user_id: User ID

        Returns:
            {
                "documents": {doc_id: markdown_content},
                "document_names": {doc_id: doc_name},
                "document_versions": {doc_id: materialization_revision},
                "failed": [doc_id],
                "failure_reasons": {doc_id: reason_code}
            }
        """
        # Verify access permission (admin users can access any KB)
        await self._verify_kb_access(kb_id, user_id)

        requested_ids: list[str] = []
        query_ids: list[uuid.UUID] = []
        seen_ids: set[str] = set()
        failure_reasons: dict[str, str] = {}
        for raw_doc_id in doc_ids:
            raw_value = str(raw_doc_id).strip()
            try:
                parsed_id = uuid.UUID(raw_value)
                normalized_id = str(parsed_id)
            except (ValueError, AttributeError):
                normalized_id = raw_value
                parsed_id = None
            if normalized_id in seen_ids:
                continue
            seen_ids.add(normalized_id)
            requested_ids.append(normalized_id)
            if parsed_id is not None:
                query_ids.append(parsed_id)
            else:
                failure_reasons[normalized_id] = "invalid_id"

        if not requested_ids:
            return {
                "documents": {},
                "document_names": {},
                "document_versions": {},
                "failed": [],
                "failure_reasons": {},
            }

        snapshots: list[tuple[str, str, str, str, str | None]] = []
        found_ids: set[str] = set()
        if query_ids:
            result = await self.db.execute(
                select(
                    Document.id,
                    Document.name,
                    Document.markdown_path,
                    Document.materialization_revision,
                    Document.markdown_sha256,
                ).where(
                    Document.kb_id == uuid.UUID(str(kb_id)),
                    Document.id.in_(query_ids),
                )
            )
            for row in result.all():
                doc_id = str(row.id)
                found_ids.add(doc_id)
                markdown_path = str(row.markdown_path or "").strip()
                if not markdown_path:
                    failure_reasons[doc_id] = "missing_markdown"
                    continue
                revision = int(row.materialization_revision or 0)
                if revision < 1:
                    failure_reasons[doc_id] = "missing_version"
                    continue
                expected_sha256 = str(row.markdown_sha256 or "").strip().lower()
                if (
                    expected_sha256
                    and re.fullmatch(r"[0-9a-f]{64}", expected_sha256) is None
                ):
                    failure_reasons[doc_id] = "invalid_hash"
                    continue
                snapshots.append(
                    (
                        doc_id,
                        str(row.name),
                        markdown_path,
                        str(revision),
                        expected_sha256 or None,
                    )
                )

        for requested_id in requested_ids:
            if requested_id not in failure_reasons and requested_id not in found_ids:
                failure_reasons[requested_id] = "not_found"

        logger.info(
            "Batch loading %s materialized Markdown document(s) out of %s requested in KB %s",
            len(snapshots),
            len(requested_ids),
            kb_id,
        )

        semaphore = asyncio.Semaphore(self.MARKDOWN_DOWNLOAD_CONCURRENCY)

        async def load_single_doc(
            snapshot: tuple[str, str, str, str, str | None],
        ) -> tuple[str, str | None, str, str, str | None]:
            doc_id, doc_name, markdown_path, version, expected_sha256 = snapshot
            object_name = markdown_path.removeprefix(f"{settings.MINIO_BUCKET}/")
            try:
                async with semaphore:
                    async with temporary_download(
                        object_name,
                        suffix=".md",
                        max_bytes=settings.MAX_UPLOAD_SIZE,
                    ) as temp_path:
                        try:
                            markdown_content = await asyncio.to_thread(
                                self._read_text_file,
                                temp_path,
                            )
                        except UnicodeError as exc:
                            _log_document_failure(
                                stage="markdown_decode",
                                document_id=doc_id,
                                exc=exc,
                            )
                            return doc_id, None, doc_name, version, "decode_error"
                if not markdown_content.strip():
                    return doc_id, None, doc_name, version, "empty_markdown"
                actual_sha256 = hashlib.sha256(
                    markdown_content.encode("utf-8")
                ).hexdigest()
                if expected_sha256 is not None and actual_sha256 != expected_sha256:
                    logger.error(
                        "document_processing stage=markdown_integrity document_id=%s task_id=none error_type=IntegrityMismatch",
                        _opaque_log_id(doc_id),
                    )
                    return doc_id, None, doc_name, version, "integrity_mismatch"
                logger.info(
                    "Loaded markdown for doc %s (%s chars)",
                    doc_id,
                    len(markdown_content),
                )
                return doc_id, markdown_content, doc_name, version, None
            except Exception as exc:
                _log_document_failure(
                    stage="markdown_download",
                    document_id=doc_id,
                    exc=exc,
                )
                return doc_id, None, doc_name, version, "storage_error"

        results = await asyncio.gather(
            *(load_single_doc(snapshot) for snapshot in snapshots)
        )

        documents: dict[str, str] = {}
        document_names: dict[str, str] = {}
        document_versions: dict[str, str] = {}
        for doc_id, content, doc_name, version, failure_reason in results:
            if content is None:
                failure_reasons[doc_id] = failure_reason or "storage_error"
                continue
            documents[doc_id] = content
            document_names[doc_id] = doc_name
            document_versions[doc_id] = version
            failure_reasons.pop(doc_id, None)

        failed = [doc_id for doc_id in requested_ids if doc_id not in documents]
        logger.info(
            "Batch load complete: %s succeeded, %s failed",
            len(documents),
            len(failed),
        )

        return {
            "documents": documents,
            "document_names": document_names,
            "document_versions": document_versions,
            "failed": failed,
            "failure_reasons": {doc_id: failure_reasons[doc_id] for doc_id in failed},
        }

    async def retry_document(
        self,
        doc_id: str,
        kb_id: str,
        user_id: str,
    ) -> dict:
        """
        Retry processing a failed document.
        Only owner and admin users can retry documents.

        Smart retry logic:
        - If markdown already exists (e.g., MinerU conversion succeeded but chunking failed),
          skip the conversion step and only retry chunking/embedding
        - Otherwise, re-run the full pipeline

        Note: Before retry, we clean up any existing ES data to avoid duplicates.
        """
        # Verify KB write access (owner or admin only)
        await self._verify_kb_write_access(kb_id, user_id)

        doc = await self.doc_repo.get_by_id(doc_id, kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={
                    "error": {"code": "NOT_FOUND", "message": "Document not found"}
                },
            )

        # Only allow retry for failed documents
        if doc.status != Document.STATUS_FAILED:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "INVALID_STATUS",
                        "message": f"Can only retry failed documents. Current status: {doc.status}",
                    }
                },
            )

        # Check if original file exists
        if not doc.file_path:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "FILE_NOT_FOUND",
                        "message": "Original file not found, please re-upload",
                    }
                },
            )

        # Persist the recoverable state before attempting the Redis write. A
        # periodic reconciler will enqueue this record if Redis is unavailable.
        await self.doc_repo.update_status(
            doc,
            Document.STATUS_QUEUED,
            error_message=None,
        )
        await self._enqueue_persisted_document(str(doc.id))
        logger.info("document_processing stage=retry_queued document_id=%s", doc_id)

        return {
            "id": str(doc.id),
            "name": doc.name,
            "status": Document.STATUS_PROCESSING,
        }

    async def _retry_chunking_only(
        self, doc_id: str, es_index_name: str, markdown_content: str, filename: str
    ):
        """
        Retry only the chunking/embedding step when markdown already exists.
        This is much faster than re-running MinerU conversion.
        """
        from config.database import AsyncSessionLocal
        from modules.knowledge.repositories.document_repository import (
            DocumentRepository,
        )
        from modules.knowledge.repositories.kb_repository import KnowledgeBaseRepository

        async with AsyncSessionLocal() as db:
            doc_repo = DocumentRepository(db)
            kb_repo = KnowledgeBaseRepository(db)

            result = await db.execute(select(Document).where(Document.id == doc_id))
            doc = result.scalar_one_or_none()
            if not doc:
                logger.error(
                    "document_processing stage=retry_load document_id=%s task_id=none error_type=DocumentNotFound",
                    _opaque_log_id(doc_id),
                )
                return

            logger.info(
                "document_processing stage=index_retry document_id=%s",
                _opaque_log_id(doc_id),
            )

            try:
                # Go directly to chunking step
                await doc_repo.update_status(doc, Document.STATUS_CHUNKING)

                parse_filename = os.path.splitext(filename)[0] + ".md"
                with self._temporary_markdown_file(
                    markdown_content, doc_id
                ) as temp_file_path:
                    parse_result = await DocumentProcessService.parse_document(
                        str(temp_file_path), str(doc_id), es_index_name, parse_filename
                    )

                task_id = parse_result["task_id"]
                logger.info(
                    "document_processing stage=index_accepted document_id=%s task_id=%s",
                    _opaque_log_id(doc_id),
                    _opaque_log_id(task_id),
                )

                await doc_repo.update_status(
                    doc, Document.STATUS_EMBEDDING, parse_task_id=task_id
                )

                # Poll parsing status
                success = await self._poll_parse_task(doc, task_id, doc_repo)

                if success:
                    logger.info(
                        "document_processing stage=complete document_id=%s task_id=%s",
                        _opaque_log_id(doc_id),
                        _opaque_log_id(task_id),
                    )
                    await kb_repo.sync_contents_count(doc.kb_id)

            except Exception as exc:
                stage = (
                    exc.stage
                    if isinstance(exc, DocumentProcessingError)
                    else "index_submit"
                )
                _log_document_failure(
                    stage=stage,
                    document_id=doc_id,
                    exc=exc,
                )
                await doc_repo.update_status(
                    doc,
                    Document.STATUS_FAILED,
                    error_message=public_document_error(
                        exc,
                        default_stage="index_submit",
                    ),
                )

    async def delete_document(self, doc_id: str, kb_id: str, user_id: str):
        """
        Delete document from KB, MinIO, and ES.
        Only owner and admin users can delete documents.
        """
        # Verify KB write access (owner or admin only)
        kb = await self._verify_kb_write_access(kb_id, user_id)

        doc = await self.doc_repo.get_by_id(doc_id, kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={
                    "error": {"code": "NOT_FOUND", "message": "Document not found"}
                },
            )

        if doc.status in {
            Document.STATUS_QUEUED,
            Document.STATUS_PROCESSING,
            Document.STATUS_CHUNKING,
            Document.STATUS_EMBEDDING,
        }:
            try:
                cancellation_settled = await cancel_document_task(str(doc.id))
            except Exception as exc:
                _log_document_failure(
                    stage="queue_cancel",
                    document_id=doc.id,
                    exc=exc,
                )
                raise HTTPException(
                    status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                    detail={
                        "error": {
                            "code": "QUEUE_UNAVAILABLE",
                            "message": "Document processing could not be stopped; retry deletion shortly",
                        }
                    },
                ) from None
            if not cancellation_settled:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail={
                        "error": {
                            "code": "PROCESSING_CANCELLATION_PENDING",
                            "message": "Document processing is still stopping; retry deletion shortly",
                        }
                    },
                )
            await self.db.refresh(doc)

        if doc.status == Document.STATUS_EMBEDDING and doc.parse_task_id:
            try:
                rag_task_settled = await self._settle_rag_task_for_deletion(
                    doc.parse_task_id,
                )
            except Exception as exc:
                _log_document_failure(
                    stage="cancellation",
                    document_id=doc.id,
                    task_id=doc.parse_task_id,
                    exc=exc,
                )
                raise HTTPException(
                    status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                    detail={
                        "error": {
                            "code": "RAG_UNAVAILABLE",
                            "message": "Document processing cancellation could not be confirmed",
                        }
                    },
                ) from None
            if not rag_task_settled:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail={
                        "error": {
                            "code": "RAG_CANCELLATION_PENDING",
                            "message": "Document processing is still stopping; retry deletion shortly",
                        }
                    },
                )

        # Delete from the KB owner's ES index to match the original write location.
        owner_es_index = self._get_owner_es_index(kb)

        # Active jobs may already have written partial chunks, so cleanup is
        # required regardless of the last database status.
        try:
            await DocumentProcessService.delete_document_from_es(
                str(doc.id), owner_es_index
            )
        except Exception as exc:
            _log_document_failure(
                stage="search_delete",
                document_id=doc.id,
                exc=exc,
                level=logging.WARNING,
            )

        # Delete from MinIO (original file)
        if doc.file_path:
            try:
                object_name = doc.file_path.replace(f"{settings.MINIO_BUCKET}/", "")
                await delete_file(object_name)
            except Exception as exc:
                _log_document_failure(
                    stage="source_delete",
                    document_id=doc.id,
                    exc=exc,
                    level=logging.WARNING,
                )

        # Delete markdown from MinIO
        if doc.markdown_path:
            try:
                md_object_name = doc.markdown_path.replace(
                    f"{settings.MINIO_BUCKET}/", ""
                )
                await delete_file(md_object_name)
            except Exception as exc:
                _log_document_failure(
                    stage="markdown_delete",
                    document_id=doc.id,
                    exc=exc,
                    level=logging.WARNING,
                )

        # Delete from DB
        await self.doc_repo.delete(doc)

        await self.kb_repo.sync_contents_count(kb_id)
        logger.info("Deleted document %s and synchronized contents count", doc_id)

    async def move_document(
        self, doc_id: str, source_kb_id: str, target_kb_id: str, user_id: str
    ) -> dict:
        """
        Move document from one knowledge base to another.
        Only owner can move documents between their own knowledge bases.

        Note: Since ES index is user-level (not KB-level), no ES migration needed.
        MinIO file paths also don't need to change as they're stored by user_id.
        """
        # Verify write access to both KBs first.
        source_kb = await self._verify_kb_write_access(source_kb_id, user_id)
        target_kb = await self._verify_kb_write_access(target_kb_id, user_id)

        # Moving documents is only safe within the same owner's KB namespace.
        if str(source_kb.owner_id) != user_id or str(target_kb.owner_id) != user_id:
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail={
                    "error": {
                        "code": "FORBIDDEN",
                        "message": "Only the owner can move documents between their own knowledge bases",
                    }
                },
            )
        if source_kb.owner_id != target_kb.owner_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "INVALID_REQUEST",
                        "message": "Documents can only be moved between knowledge bases owned by the same user",
                    }
                },
            )

        # Get document
        doc = await self.doc_repo.get_by_id(doc_id, source_kb_id)
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail={
                    "error": {"code": "NOT_FOUND", "message": "Document not found"}
                },
            )

        # Cannot move to the same KB
        if source_kb_id == target_kb_id:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail={
                    "error": {
                        "code": "INVALID_REQUEST",
                        "message": "Source and target knowledge base are the same",
                    }
                },
            )

        # Move document (update kb_id)
        await self.doc_repo.update_kb_id(doc, target_kb_id)

        # Update contents count for both KBs (only for ready documents)
        if doc.status == Document.STATUS_READY:
            await self.kb_repo.increment_contents_count(source_kb_id, -1)
            await self.kb_repo.increment_contents_count(target_kb_id, 1)
            logger.info(
                f"Moved document {doc_id} from KB {source_kb_id} to KB {target_kb_id}"
            )
        else:
            logger.info(
                f"Moved document {doc_id} (status: {doc.status}, no count change)"
            )

        return {
            "id": str(doc.id),
            "name": doc.name,
            "sourceKbId": source_kb_id,
            "targetKbId": target_kb_id,
            "status": doc.status,
        }
