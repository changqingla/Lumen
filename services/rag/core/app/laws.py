#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

from tika import parser
import re
from io import BytesIO
from docx import Document

from deepdoc.parser.utils import get_text
from core.nlp import bullets_category, remove_contents_table, hierarchical_merge, \
    make_colon_as_title, tokenize_chunks, docx_question_level
from core.nlp import rag_tokenizer
from deepdoc.parser import DocxParser, HtmlParser


class Docx(DocxParser):
    def __init__(self):
        pass

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        bull = bullets_category([p.text for p in self.doc.paragraphs])
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = docx_question_level(p, bull)
            if not p_text.strip("\n"):
                continue
            lines.append((question_level, p_text))

            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1

        visit = [False for _ in range(len(lines))]
        sections = []
        for s in range(len(lines)):
            e = s + 1
            while e < len(lines):
                if lines[e][0] <= lines[s][0]:
                    break
                e += 1
            if e - s == 1 and visit[s]:
                continue
            sec = []
            next_level = lines[s][0] + 1
            while not sec and next_level < 22:
                for i in range(s+1, e):
                    if lines[i][0] != next_level:
                        continue
                    sec.append(lines[i][1])
                    visit[i] = True
                next_level += 1
            sec.insert(0, lines[s][1])

            sections.append("\n".join(sec))
        return [s for s in sections if s]

    def __str__(self) -> str:
        return f'''
            question:{self.question},
            answer:{self.answer},
            level:{self.level},
            childs:{self.childs}
        '''


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
        Supported file formats are docx, txt.
    """
    doc = {
        "docnm_kwd": filename,
        "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))
    }
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    sections = []
    # is it English
    eng = lang.lower() == "english"

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunks = Docx()(filename, binary)
        callback(0.7, "Finish parsing.")
        return tokenize_chunks(chunks, doc, eng)

    elif re.search(r"\.txt$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        sections = txt.split("\n")
        sections = [s for s in sections if s]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = HtmlParser()(filename, binary)
        sections = [s for s in sections if s]
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        binary = BytesIO(binary)
        doc_parsed = parser.from_buffer(binary)
        sections = doc_parsed['content'].split('\n')
        sections = [s for s in sections if s]
        callback(0.8, "Finish parsing.")

    else:
        raise NotImplementedError(
            "file type not supported yet(doc, docx, txt supported)")


    # Remove 'Contents' part
    remove_contents_table(sections, eng)

    make_colon_as_title(sections)
    bull = bullets_category(sections)
    chunks = hierarchical_merge(bull, sections, 5)
    if not chunks:
        callback(0.99, "No chunk parsed out.")

    return tokenize_chunks(["\n".join(ck)
                           for ck in chunks], doc, eng)
