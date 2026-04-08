#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

import logging
import copy
import re

from io import BytesIO
from core.nlp import rag_tokenizer, tokenize, tokenize_table, bullets_category, title_frequency, tokenize_chunks, docx_question_level
from core.utils import num_tokens_from_string
from deepdoc.parser import DocxParser
from docx import Document
from PIL import Image


class Docx(DocxParser):
    def __init__(self):
        pass

    def get_picture(self, document, paragraph):
        img = paragraph._element.xpath('.//pic:pic')
        if not img:
            return None
        img = img[0]
        embed = img.xpath('.//a:blip/@r:embed')[0]
        try:
            related_part = document.part.related_parts[embed]
        except KeyError:
            # 跳过损坏的图片引用（如指向 NULL 的关系）
            logging.warning(f"跳过损坏的图片引用: {embed}")
            return None
        try:
            image = related_part.image
            image = Image.open(BytesIO(image.blob))
            return image
        except (AttributeError, Exception) as e:
            logging.warning(f"无法读取图片 {embed}: {e}")
            return None

    def concat_img(self, img1, img2):
        if img1 and not img2:
            return img1
        if not img1 and img2:
            return img2
        if not img1 and not img2:
            return None
        width1, height1 = img1.size
        width2, height2 = img2.size

        new_width = max(width1, width2)
        new_height = height1 + height2
        new_image = Image.new('RGB', (new_width, new_height))

        new_image.paste(img1, (0, 0))
        new_image.paste(img2, (0, height1))

        return new_image

    def __call__(self, filename, binary=None, from_page=0, to_page=100000, callback=None):
        self.doc = Document(
            filename) if not binary else Document(BytesIO(binary))
        pn = 0
        last_answer, last_image = "", None
        question_stack, level_stack = [], []
        ti_list = []
        for p in self.doc.paragraphs:
            if pn > to_page:
                break
            question_level, p_text = 0, ''
            if from_page <= pn < to_page and p.text.strip():
                question_level, p_text = docx_question_level(p)
            if not question_level or question_level > 6: # not a question
                last_answer = f'{last_answer}\n{p_text}'
                current_image = self.get_picture(self.doc, p)
                last_image = self.concat_img(last_image, current_image)
            else:   # is a question
                if last_answer or last_image:
                    sum_question = '\n'.join(question_stack)
                    if sum_question:
                        ti_list.append((f'{sum_question}\n{last_answer}', last_image))
                    last_answer, last_image = '', None

                i = question_level
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(p_text)
                level_stack.append(question_level)
            for run in p.runs:
                if 'lastRenderedPageBreak' in run._element.xml:
                    pn += 1
                    continue
                if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        if last_answer:
            sum_question = '\n'.join(question_stack)
            if sum_question:
                ti_list.append((f'{sum_question}\n{last_answer}', last_image))
                
        tbls = []
        for tb in self.doc.tables:
            html= "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    for j in range(i+1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))
        return ti_list, tbls


def chunk(filename, binary=None, from_page=0, to_page=100000,
          lang="Chinese", callback=None, **kwargs):
    """
        Only docx is supported.
    """
    doc = {
        "docnm_kwd": filename
    }
    doc["title_tks"] = rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", doc["docnm_kwd"]))
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    # is it English
    eng = lang.lower() == "english"

    if re.search(r"\.docx?$", filename, re.IGNORECASE):
        docx_parser = Docx()
        ti_list, tbls = docx_parser(filename, binary,
                                    from_page=0, to_page=10000, callback=callback)
        res = tokenize_table(tbls, doc, eng)
        for text, image in ti_list:
            d = copy.deepcopy(doc)
            d['image'] = image
            tokenize(d, text, eng)
            res.append(d)
        return res
    else:
        raise NotImplementedError("file type not supported yet(docx supported)")
    

if __name__ == "__main__":
    import sys


    def dummy(prog=None, msg=""):
        pass


    chunk(sys.argv[1], callback=dummy)
