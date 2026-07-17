#
#  Copyright 2025 The InfiniFlow Authors. All Rights Reserved.
#
#  Licensed under the Apache License, Version 2.0 (the "License");
#  you may not use this file except in compliance with the License.
#  You may obtain a copy of the License at
#
#      http://www.apache.org/licenses/LICENSE-2.0
#
#  Unless required by applicable law or agreed to in writing, software
#  distributed under the License is distributed on an "AS IS" BASIS,
#  WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
#  See the License for the specific language governing permissions and
#  limitations under the License.
#

from docx import Document
import re
import pandas as pd
from collections import Counter
from core.nlp import rag_tokenizer
from io import BytesIO
import logging


def _patch_docx_null_references():
    """
    修复 python-docx 库无法处理包含 NULL 引用的 docx 文件的问题。
    
    某些 Word 文档内部可能存在指向 'NULL' 的损坏关系链接，
    这会导致 python-docx 在解析时抛出 "There is no item named 'NULL' in the archive" 错误。
    
    此补丁包含两部分：
    1. 替换 _SerializedRelationships.load_from_xml 方法，在加载关系时跳过目标为 'NULL' 的损坏链接
    2. 替换 RelationshipCollection 的关系查找方法，使其在找不到关系时返回 None 而不是抛出异常
    """
    # 记录被跳过的关系ID，供后续查找时使用
    _skipped_rel_ids = set()
    
    try:
        from docx.opc.pkgreader import _SerializedRelationships, _SerializedRelationship
        from docx.opc.oxml import parse_xml
        
        @staticmethod
        def _patched_load_from_xml(baseURI, rels_item_xml):
            """修复后的 load_from_xml 方法，跳过 NULL 引用"""
            srels = _SerializedRelationships()
            if rels_item_xml is not None:
                rels_elm = parse_xml(rels_item_xml)
                for rel_elm in rels_elm.Relationship_lst:
                    # 跳过目标为 NULL 的损坏关系链接
                    target = rel_elm.target_ref
                    if target in ('NULL', '../NULL', './NULL', '/NULL'):
                        logging.warning(
                            "DOCX relation skipped: stage=load_relationships reason=null_target"
                        )
                        _skipped_rel_ids.add(rel_elm.rId)
                        continue
                    srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
            return srels
        
        # 替换原始方法
        _SerializedRelationships.load_from_xml = _patched_load_from_xml
        logging.debug("已应用 python-docx NULL 引用修复补丁 (load_from_xml)")
        
    except Exception as error:
        logging.warning(
            "DOCX compatibility patch failed: stage=load_from_xml error_type=%s",
            type(error).__name__,
        )
    
    # 补丁2：修复 RelationshipCollection 的 get_or_add 和 __getitem__ 方法
    try:
        from docx.opc.rel import Relationships
        
        _original_getitem = Relationships.__getitem__
        
        def _patched_getitem(self, rId):
            """修复后的 __getitem__ 方法，对跳过的关系ID返回 None"""
            try:
                return _original_getitem(self, rId)
            except KeyError:
                if rId in _skipped_rel_ids:
                    logging.debug("DOCX skipped relation lookup ignored")
                    return None
                raise
        
        Relationships.__getitem__ = _patched_getitem
        logging.debug("已应用 python-docx NULL 引用修复补丁 (Relationships.__getitem__)")
        
    except Exception as error:
        logging.warning(
            "DOCX compatibility patch failed: stage=relationships error_type=%s",
            type(error).__name__,
        )
    
    # 补丁3：修复 Part 的 related_parts 属性访问
    try:
        from docx.opc.part import Part
        
        if hasattr(Part, 'related_part'):
            _original_related_part = Part.related_part
            
            def _patched_related_part(self, rId):
                """修复后的 related_part 方法，对损坏的关系返回 None"""
                try:
                    return _original_related_part(self, rId)
                except (KeyError, AttributeError):
                    if rId in _skipped_rel_ids:
                        logging.debug("DOCX damaged related-part lookup ignored")
                        return None
                    raise
            
            Part.related_part = _patched_related_part
            logging.debug("已应用 python-docx NULL 引用修复补丁 (Part.related_part)")
            
    except Exception as error:
        logging.warning(
            "DOCX compatibility patch failed: stage=related_part error_type=%s",
            type(error).__name__,
        )


# 在模块加载时应用补丁
_patch_docx_null_references()


class DeepRAGDocxParser:

    def __extract_table_content(self, tb):
        df = []
        for row in tb.rows:
            df.append([c.text for c in row.cells])
        return self.__compose_table_content(pd.DataFrame(df))

    def __compose_table_content(self, df):

        def blockType(b):
            patt = [
                ("^(20|19)[0-9]{2}[年/-][0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^(20|19)[0-9]{2}年$", "Dt"),
                (r"^(20|19)[0-9]{2}[年/-][0-9]{1,2}月*$", "Dt"),
                ("^[0-9]{1,2}[月/-][0-9]{1,2}日*$", "Dt"),
                (r"^第*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}年*[一二三四1-4]季度$", "Dt"),
                (r"^(20|19)[0-9]{2}[ABCDE]$", "DT"),
                ("^[0-9.,+%/ -]+$", "Nu"),
                (r"^[0-9A-Z/\._~-]+$", "Ca"),
                (r"^[A-Z]*[a-z' -]+$", "En"),
                (r"^[0-9.,+-]+[0-9A-Za-z/$￥%<>（）()' -]+$", "NE"),
                (r"^.{1}$", "Sg")
            ]
            for p, n in patt:
                if re.search(p, b):
                    return n
            tks = [t for t in rag_tokenizer.tokenize(b).split() if len(t) > 1]
            if len(tks) > 3:
                if len(tks) < 12:
                    return "Tx"
                else:
                    return "Lx"

            if len(tks) == 1 and rag_tokenizer.tag(tks[0]) == "nr":
                return "Nr"

            return "Ot"

        if len(df) < 2:
            return []
        max_type = Counter([blockType(str(df.iloc[i, j])) for i in range(
            1, len(df)) for j in range(len(df.iloc[i, :]))])
        max_type = max(max_type.items(), key=lambda x: x[1])[0]

        colnm = len(df.iloc[0, :])
        hdrows = [0]  # header is not nessesarily appear in the first line
        if max_type == "Nu":
            for r in range(1, len(df)):
                tys = Counter([blockType(str(df.iloc[r, j]))
                              for j in range(len(df.iloc[r, :]))])
                tys = max(tys.items(), key=lambda x: x[1])[0]
                if tys != max_type:
                    hdrows.append(r)

        lines = []
        for i in range(1, len(df)):
            if i in hdrows:
                continue
            hr = [r - i for r in hdrows]
            hr = [r for r in hr if r < 0]
            t = len(hr) - 1
            while t > 0:
                if hr[t] - hr[t - 1] > 1:
                    hr = hr[t:]
                    break
                t -= 1
            headers = []
            for j in range(len(df.iloc[i, :])):
                t = []
                for h in hr:
                    x = str(df.iloc[i + h, j]).strip()
                    if x in t:
                        continue
                    t.append(x)
                t = ",".join(t)
                if t:
                    t += ": "
                headers.append(t)
            cells = []
            for j in range(len(df.iloc[i, :])):
                if not str(df.iloc[i, j]):
                    continue
                cells.append(headers[j] + str(df.iloc[i, j]))
            lines.append(";".join(cells))

        if colnm > 3:
            return lines
        return ["\n".join(lines)]

    def __call__(self, fnm, from_page=0, to_page=100000000):
        self.doc = Document(fnm) if isinstance(
            fnm, str) else Document(BytesIO(fnm))
        pn = 0 # parsed page
        secs = [] # parsed contents
        
        try:
            for p in self.doc.paragraphs:
                if pn > to_page:
                    break

                runs_within_single_paragraph = [] # save runs within the range of pages
                try:
                    for run in p.runs:
                        if pn > to_page:
                            break
                        if from_page <= pn < to_page and p.text.strip():
                            runs_within_single_paragraph.append(run.text) # append run.text first

                        # wrap page break checker into a static method
                        try:
                            if 'lastRenderedPageBreak' in run._element.xml:
                                pn += 1
                        except Exception:
                            pass
                except KeyError as error:
                    # 跳过因损坏关系导致的错误
                    logging.warning(
                        "DOCX content skipped: stage=paragraph_run error_type=%s",
                        type(error).__name__,
                    )
                    continue

                secs.append(("".join(runs_within_single_paragraph), p.style.name if hasattr(p.style, 'name') else '')) # then concat run.text as part of the paragraph
        except KeyError as error:
            logging.warning(
                "DOCX content parsing stopped: stage=paragraphs error_type=%s",
                type(error).__name__,
            )

        tbls = []
        try:
            for tb in self.doc.tables:
                try:
                    tbls.append(self.__extract_table_content(tb))
                except KeyError as error:
                    logging.warning(
                        "DOCX content skipped: stage=table error_type=%s",
                        type(error).__name__,
                    )
                    continue
        except KeyError as error:
            logging.warning(
                "DOCX content parsing stopped: stage=tables error_type=%s",
                type(error).__name__,
            )
            
        return secs, tbls
